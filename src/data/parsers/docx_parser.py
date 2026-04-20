"""
AegisRAG - DOCX Parser

Extracts paragraph-level text blocks from Microsoft Word ``.docx`` files
using the ``python-docx`` library. Tables are flattened to tab/newline
separated text and emitted as additional blocks.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse a DOCX file into a list of paragraph-level text blocks.

    Output schema: ``[{"text": str, "page_number": None}, ...]``.
    Page numbers are not recoverable from .docx; the key is kept for
    uniformity with :class:`PDFParser`.
    """

    def __init__(self) -> None:
        try:
            import docx  # noqa: F401
        except ImportError as exc:
            raise ImportError(
                "DOCXParser requires 'python-docx'. "
                "Install with: pip install python-docx"
            ) from exc

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(self, path: Path) -> list[dict[str, Any]]:
        """Parse a DOCX file and return paragraph + table text blocks.

        Parameters
        ----------
        path : Path
            Path to the ``.docx`` file.

        Returns
        -------
        list of dict
            Each dict is ``{"text": str, "page_number": None}``.
        """
        import docx

        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")
        if not path.is_file():
            raise ValueError(f"Not a file: {path}")

        logger.info("Parsing DOCX: %s", path)
        document = docx.Document(str(path))

        blocks: list[dict[str, Any]] = []

        # Paragraphs
        for para in document.paragraphs:
            text = (para.text or "").strip()
            if text:
                blocks.append({"text": text, "page_number": None})

        # Tables flattened
        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [(cell.text or "").strip() for cell in row.cells]
                rows.append("\t".join(cells))
            flat = "\n".join(rows).strip()
            if flat:
                blocks.append(
                    {"text": f"[TABLE]\n{flat}", "page_number": None}
                )

        logger.info(
            "DOCX parsed: %d blocks extracted from %s", len(blocks), path.name
        )
        return blocks
